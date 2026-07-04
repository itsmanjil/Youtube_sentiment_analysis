#!/usr/bin/env python3
"""
Compile the thesis-facing Markdown chapters into a single formatted thesis.docx.

Renders headings, paragraphs, bullet/numbered lists, GitHub-style tables,
inline **bold** and `code`, and block code fences into a Word document with a
title page and an automatic Table of Contents field.

Usage
-----
    cd backend
    python scripts/build_thesis_docx.py --out ../thesis.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

BACKEND = Path(__file__).resolve().parents[1]
DOCS = BACKEND / "docs"

# (heading shown in thesis, source markdown file, drop the file's own H1 title?)
CHAPTERS = [
    ("Chapter 1  Introduction and Research Questions", "THESIS_CHAPTER1_RESEARCH_QUESTIONS.md", True),
    ("Chapter 2  Literature Review",                   "THESIS_LITERATURE_REVIEW.md",          True),
    ("Chapter 3  Methodology and Data",                "LABEL_PROVENANCE.md",                  True),
    ("Chapter 3 (continued)  Exploratory Data Analysis", "THESIS_EDA.md",                      True),
    ("Chapter 3B  System Design and Implementation",   "THESIS_CHAPTER3B_SYSTEM_DESIGN.md",    True),
    ("Chapter 4  Consolidated Evaluation",             "THESIS_EVALUATION_CONSOLIDATED.md",    True),
    ("Chapter 4 (Discussion)  Results Narrative",      "THESIS_ABSTRACT_CHAPTER4_POLISHED.md", True),
    ("Chapter 5  Conclusion and Future Work",          "THESIS_CHAPTER5_CONCLUSION_FUTURE_WORK.md", True),
    ("Threats to Validity",                            "THESIS_RISKS_GAPS.md",                 True),
    ("Appendix A  Claim-to-Artifact Audit",            "THESIS_CLAIM_ARTIFACT_AUDIT.md",       True),
]

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_runs(paragraph, text: str) -> None:
    """Add a paragraph's text, honouring **bold** and `code` inline spans."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_toc(document: Document) -> None:
    """Insert an auto-updating Table of Contents field."""
    p = document.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose 'Update Field' to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar, instr, fldChar2, t, fldChar3):
        run._r.append(el)


def parse_table(lines: List[str], i: int, document: Document) -> int:
    """Render a contiguous block of markdown table rows starting at line i."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    # rows[1] is the |---|---| separator
    parsed = []
    for idx, row in enumerate(rows):
        if idx == 1:
            continue
        cells = [c.strip() for c in row.strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return i
    ncols = max(len(r) for r in parsed)
    table = document.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ridx, cells in enumerate(parsed):
        cells = cells + [""] * (ncols - len(cells))
        wrow = table.add_row().cells
        for cidx, cell in enumerate(cells):
            para = wrow[cidx].paragraphs[0]
            add_runs(para, cell.replace("\\|", "|"))
            if ridx == 0:
                for run in para.runs:
                    run.bold = True
    document.add_paragraph()
    return i


def render_markdown(md: str, document: Document, drop_h1: bool) -> None:
    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf: List[str] = []
    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            if in_code:
                p = document.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line); i += 1; continue

        stripped = line.strip()

        # tables
        if stripped.startswith("|"):
            i = parse_table(lines, i, document); continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1)); text = m.group(2).strip()
            if level == 1 and drop_h1:
                i += 1; continue
            document.add_heading(text, level=min(level + 1, 4) if not drop_h1 else level)
            i += 1; continue

        # blank
        if not stripped:
            i += 1; continue

        # horizontal rule
        if stripped in ("---", "***", "___"):
            i += 1; continue

        # bullet list
        bm = re.match(r"^[-*]\s+(.*)$", stripped)
        if bm:
            p = document.add_paragraph(style="List Bullet"); add_runs(p, bm.group(1)); i += 1; continue

        # numbered list
        nm = re.match(r"^\d+\.\s+(.*)$", stripped)
        if nm:
            p = document.add_paragraph(style="List Number"); add_runs(p, nm.group(1)); i += 1; continue

        # blockquote
        qm = re.match(r"^>\s?(.*)$", stripped)
        if qm:
            p = document.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(qm.group(1)); r.italic = True; i += 1; continue

        # normal paragraph
        p = document.add_paragraph(); add_runs(p, stripped); i += 1


def build_title_page(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    t = document.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Uncertainty-Aware Ensemble Sentiment Analysis\nfor YouTube Comments")
    r.bold = True; r.font.size = Pt(22)
    s = document.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("A Computational Intelligence Approach with Multi-Objective\n"
                   "Optimisation, Calibration, and Selective Prediction")
    rs.italic = True; rs.font.size = Pt(14)
    for _ in range(6):
        document.add_paragraph()
    deg = document.add_paragraph(); deg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    deg.add_run("MSc Data Science and Computational Intelligence").font.size = Pt(12)
    document.add_page_break()


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=str(BACKEND.parent / "thesis.docx"))
    args = ap.parse_args()

    document = Document()
    # base style
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)

    build_title_page(document)

    # Abstract (pulled from the polished draft's Abstract section)
    abstract_src = (DOCS / "THESIS_ABSTRACT_CHAPTER4_POLISHED.md").read_text(encoding="utf-8")
    abs_match = re.search(r"## Abstract\s+(.*?)(?=\n## )", abstract_src, re.S)
    document.add_heading("Abstract", level=1)
    if abs_match:
        for para in abs_match.group(1).strip().split("\n\n"):
            p = document.add_paragraph(); add_runs(p, para.replace("\n", " ").strip())
    document.add_page_break()

    # Table of contents
    document.add_heading("Table of Contents", level=1)
    add_toc(document)
    document.add_page_break()

    # Chapters
    for title, fname, drop_h1 in CHAPTERS:
        path = DOCS / fname
        if not path.exists():
            print(f"  WARN: missing {fname}, skipping")
            continue
        document.add_heading(title, level=1)
        md = path.read_text(encoding="utf-8")
        # strip the status-date and reproduce lines for cleanliness
        md = re.sub(r"^Status date:.*$", "", md, flags=re.M)
        md = re.sub(r"^Generated artifact:.*$", "", md, flags=re.M)
        md = re.sub(r"^Reproduce:.*$", "", md, flags=re.M)
        md = re.sub(r"^Length:.*$", "", md, flags=re.M)
        render_markdown(md, document, drop_h1)
        document.add_page_break()
        print(f"  added: {title}  <- {fname}")

    document.save(args.out)
    print(f"\nSaved {args.out}")


if __name__ == "__main__":
    main()
